from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os

def add_hyperlink(paragraph, text, url):
    # This places a hyperlink within a paragraph object.
    # Create the w:hyperlink tag
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create the w:r element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Join all the xml elements together to form the hyperlink style
    # Style: Hyperlink
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    
    # Text color blue
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1") # Standard Word hyperlink blue
    rPr.append(color)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def process_line_with_links(paragraph, line_text):
    # Regex to find Markdown links: [text](url)
    # We need to split the text by links and add runs/hyperlinks sequentially
    pattern = re.compile(r'\[(.*?)\]\((.*?)\)')
    last_pos = 0
    
    for match in pattern.finditer(line_text):
        # Add text before the link
        start, end = match.span()
        if start > last_pos:
            paragraph.add_run(line_text[last_pos:start])
        
        # Add the link
        link_text = match.group(1)
        link_url = match.group(2)
        add_hyperlink(paragraph, link_text, link_url)
        
        last_pos = end
        
    # Add remaining text
    if last_pos < len(line_text):
        paragraph.add_run(line_text[last_pos:])

def create_docx(source_txt, dest_docx):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(source_txt, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table = None # State to track active table

    for line in lines:
        line = line.strip()
        if not line:
            table = None # End of table on empty line
            continue

        # Header detection (Revised for "Section A:", etc.)
        if line.startswith("Subject:") or line.lower().startswith("section"):
            table = None
            p = doc.add_paragraph()
            runner = p.add_run(line)
            runner.bold = True
            if line.lower().startswith("section"):
                runner.font.size = Pt(12)
                p.paragraph_format.space_before = Pt(12)
        
        # Table Row Detection (Pipe delimited)
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            
            # If no active table, create one
            if table is None:
                table = doc.add_table(rows=0, cols=len(cells))
                table.style = 'Table Grid'
                table.autofit = True
            
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(cells):
                if i < len(row_cells):
                    # Handle links inside table cells too
                    process_line_with_links(row_cells[i].paragraphs[0], cell_text)
                    
                    # Bold header row if it's the first row (heuristic: matches "Challenge")
                    if cell_text == "Challenge":
                        for run in row_cells[i].paragraphs[0].runs:
                            run.bold = True

        # Separator
        elif line.startswith("====="):
            table = None
            doc.add_paragraph("_" * 30)
            
        # Metadata keys - bold keys but handle links in value
        elif any(line.startswith(k) for k in ["Name:", "Project Name:", "Date:"]):
            table = None
            p = doc.add_paragraph()
            parts = line.split(":", 1)
            p.add_run(parts[0] + ":").bold = True
            process_line_with_links(p, parts[1])
            
        # Bullet points
        elif line.startswith("- ") or line.startswith("* "):
            table = None
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:]
            # Bold keys in bullets (e.g., "**Key**:")
            if "**" in content:
                parts = content.split(":", 1)
                if len(parts) > 1:
                    bold_part = parts[0].replace("**", "")
                    p.add_run(bold_part + ":").bold = True
                    process_line_with_links(p, parts[1])
                else:
                    process_line_with_links(p, content.replace("**", ""))
            else:
                process_line_with_links(p, content)
                
        # Normal text with potential links
        else:
            table = None
            p = doc.add_paragraph()
            process_line_with_links(p, line)

    doc.save(dest_docx)
    print(f"Created: {dest_docx}")

if __name__ == "__main__":
    src = r"c:\Users\dhl\data\thesis\thesis\Thesis_Draft\Updates\2026-02-01_Status_Update_GDOC_READY.txt"
    dst = r"c:\Users\dhl\data\thesis\thesis\Thesis_Draft\Updates\2026-02-01_Status_Update.docx"
    create_docx(src, dst)
