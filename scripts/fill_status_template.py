from docx.oxml.shared import OxmlElement, qn
import re

def add_hyperlink(paragraph, text, url):
    # This places a hyperlink within a paragraph object.
    # Create the w:hyperlink tag
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create the w:r element
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    
    # Check if bold is needed (if it was passed in previous context, but here we assume standard link style)
    # Usually links are blue/underlined.
    
    # Join with "Hyperlink" style if possible, or manual formatting
    # Manual:
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1") # Default blue
    rPr.append(c)
    
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def process_line_with_links(paragraph, line_text, bold_prefix=None):
    """
    Parses line_text for Markdown links [Text](URL) and adds runs/hyperlinks to paragraph.
    If bold_prefix is provided (e.g. 'Key:'), it is added as a bold run first.
    """
    if bold_prefix:
        paragraph.add_run(bold_prefix).bold = True
    
    # Split by links
    # Pattern: [Text](URL)
    # We use capturing groups to keep separators
    parts = re.split(r'(\[.*?\]\(.*?\))', line_text)
    
    for part in parts:
        link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
        if link_match:
            text = link_match.group(1)
            url = link_match.group(2)
            add_hyperlink(paragraph, text, url)
        else:
            if part: # process potential simple urls? "https://..."
                 # Simple check for raw http links if no markdown
                 if "http" in part and not link_match:
                     # very basic split for separate urls? 
                     # For now, just add text to avoid breaking simple text
                     paragraph.add_run(part)
                 else:
                     paragraph.add_run(part)

# ... inside fill_template ...

            # Use manual bullet formatting since 'List Bullet' style is missing
            current_p = target_p
            for line in lines:
                new_p = doc.add_paragraph(style='normal') # Fallback to normal
                p_fmt = new_p.paragraph_format
                p_fmt.left_indent = Pt(18) # Indent for bullet
                p_fmt.first_line_indent = Pt(-18) # Hanging indent
                
                # Format text
                runner = new_p.add_run("• ") # Manual bullet char
                
                if "**" in line:
                    clean_line = line.replace("* ", "").replace("- ", "")
                    # Split only on the first colon to separate Key from Value
                    parts = clean_line.split(":", 1)
                    if len(parts) > 1:
                        # parts[0] is the key (e.g. "**Data Collection**")
                        # We strip ** and bold it (and add colon back)
                        key_text = parts[0].replace("**", "") + ":"
                        
                        # parts[1] is the rest. It might have markdown links.
                        process_line_with_links(new_p, parts[1], bold_prefix=key_text)
                    else:
                        # No colon.
                        process_line_with_links(new_p, clean_line.replace("**", ""), bold_prefix=None) # Maybe bold whole thing? Logic unclear, assume plain processing or passing bold flag.
                        # Actually if it was **Text**, we might want it bold. 
                        # simplicity: just process links, maybe bolding is lost for non-key items.
                else:
                    process_line_with_links(new_p, line.replace("* ", "").replace("- ", ""))
                
                # Move this new_p to right position
                new_p._element.getparent().remove(new_p._element)
                current_p._element.addnext(new_p._element)
                current_p = new_p # Advance
    """Parses the text file into metadata, sections, and table data."""
    with open(txt_path, 'r', encoding='utf-8') as f:
        lines = [l.strip() for l in f.readlines()]

    data = {
        "metadata": {},
        "sections": {"A": [], "B": [], "C": []},
        "table": [],
        "footer": []
    }
    
    current_section = None
    in_table = False
    in_footer = False

    for line in lines:
        if not line: continue
        
        # Metadata
        if line.startswith("Name:"): data["metadata"]["Name"] = line
        elif line.startswith("Project Name:"): data["metadata"]["Project"] = line
        elif line.startswith("Date:"): data["metadata"]["Date"] = line
        
        # Sections
        elif line.lower().startswith("section a:"): current_section = "A"
        elif line.lower().startswith("section b:"): current_section = "B"
        elif line.lower().startswith("section c:"): current_section = "C"
        elif line.lower().startswith("section d:"): 
            current_section = "D"
            in_table = True
        
        # Footer
        elif line.startswith("PROGRESS LINK:"):
            in_footer = True
            current_section = None
            data["footer"].append(line)
        
        elif in_footer:
            data["footer"].append(line)

        # Content Collection
        elif current_section in ["A", "B", "C"]:
            # Clean up bullet markers for inserting into doc
            data["sections"][current_section].append(line)
            
        elif current_section == "D" and in_table and line.startswith("|"):
            # Parse table row
            parts = [p.strip() for p in line.strip("|").split("|")]
            if parts[0] == "Challenge": continue # Skip header
            data["table"].append(parts)

    return data

def fill_template(template_path, output_path, data):
    doc = Document(template_path)
    
    # 1. Update Metadata (First few paragraphs)
    for p in doc.paragraphs[:10]: # Check first 10 lines
        text = p.text.strip()
        if text.startswith("Name:"):
            p.text = data["metadata"].get("Name", text)
            p.runs[0].bold = True # Re-bold attempts
        elif text.startswith("Project Name:"):
            p.text = data["metadata"].get("Project", text)
            p.runs[0].bold = True
        elif text.startswith("Date:"):
            p.text = data["metadata"].get("Date", text)
            p.runs[0].bold = True

    # 2. Update Sections A, B, C
    # Strategy: Find header, delete subsequent paragraphs until next header, insert new bullets
    
    def repl_section(sec_letter, new_lines):
        start_idx = -1
        end_idx = -1
        
        # Find start
        header_prefix = f"Section {sec_letter}:"
        next_header_prefix = "Section "
        
        paragraphs = list(doc.paragraphs) # Snapshot
        
        for i, p in enumerate(paragraphs):
            if p.text.strip().lower().startswith(header_prefix.lower()):
                start_idx = i
                continue
            
            if start_idx != -1:
                # Look for next section or end of doc (for C, next is D)
                if p.text.strip().lower().startswith(next_header_prefix.lower()) and p.text.strip().lower() != header_prefix.lower():
                    end_idx = i
                    break
                    
        # If we found the range, delete old content and insert new
        if start_idx != -1:
            # Determine range to clear. If end_idx is -1, it goes to D?
            # Actually finding D is safer.
            pass

    # Simplified Paragraph Iteration to Replace Content in-place or clear
    # Only tricky part is deleting paras while iterating.
    # New Approach: Iterate all paras, identifying ranges to DELETE, then do deletions in reverse index.
    
    paras_to_delete = []
    insert_points = {} # index -> list of lines to insert AFTER
    
    section_map = {
        "Section A:": "A",
        "Section B:": "B",
        "Section C:": "C",
        "Section D:": "D" # Stop clearing C
    }
    
    current_sec = None
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        matched_header = False
        for prefix, key in section_map.items():
            if txt.lower().startswith(prefix.lower()):
                current_sec = key
                matched_header = True
                if key in ["A", "B", "C"]:
                    insert_points[i] = data["sections"][key]
                break
        
        if not matched_header and current_sec in ["A", "B", "C"]:
            # Identify paragraphs to remove (old bullets)
            # Be careful not to delete empty spacing lines if we want to keep them, 
            # but template usually has content we want to wipe.
            if txt: # Delete non-empty lines (bullets)
                paras_to_delete.append(p)
            elif not txt: # Keep empty lines for spacing? Or delete to control it?
                pass 

    # Execute Deletions
    for p in paras_to_delete:
        p._element.getparent().remove(p._element)

    # Execute Insertions
    # We need to find the headers again because indices shifted? 
    # Actually, inserting *after* a paragraph element is stable if we have the reference.
    # But wait, insert_paragraph usually appends. to insert *after*, we need to locate matches again.
    
    for p in doc.paragraphs:
        txt = p.text.strip()
        sec_key = None
        for prefix, key in section_map.items():
            if txt.lower().startswith(prefix.lower()):
                sec_key = key
                break
        
        if sec_key and sec_key in data["sections"]:
            # Insert bullets after this header
            # We want to insert *immediately after* the header p
            # Logic: p.insert_paragraph_before() ? No, we need after.
            # docx element insert is: p._p.addnext(new_p._p)
            
            # Simple hack: Append to the header paragraph? No.
            # Iterate lines in reverse order so we can insert_paragraph_before the *next* element?
            # Or just append to document? No, need order.
            
            # Reliable way: find the parent (body) and insert index.
            pass

    # Easier approach for sections:
    # use `p.insert_paragraph_before()` on the *next* paragraph? 
    # What if D is immediately after C?
    
    # Let's re-scan paragraphs to insert.
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        for prefix, key in section_map.items():
            if txt.lower().startswith(prefix.lower()) and key in ["A", "B", "C"]:
                lines_to_add = data["sections"][key]
                # We need to insert these lines *after* p. 
                # effectively, we can insert them *before* the *next* paragraph.
                # If this is the last para, append.
                
                # Note: modifying list while iterating is bad, but we are just reading doc.paragraphs
                # actually doc.paragraphs re-evaluates?
                pass

    # OK, the python-docx `insert_paragraph_before` is on a paragraph object.
    # So if we find header, we look for the *next* paragraph and insert before it.
    
    def insert_lines_after_header(header_prefix_list, lines):
        # find header
        target_p = None
        for p in doc.paragraphs:
            for hp in header_prefix_list:
                if p.text.strip().lower().startswith(hp.lower()):
                    target_p = p
                    break
            if target_p: break
        
        if target_p:
            # We want to append paragraphs after this. 
            # finding the next sibling is the OXML way
            parent = target_p._element.getparent()
            index = parent.index(target_p._element) + 1
            
            for line in reversed(lines): # Insert in reverse so they end up in order if inserting at same index?
                # or insert at index, then index+1...
                # actually `insert_paragraph_before` logic on the *next* existing para is easier if it exists.
                pass
                
            # Let's just use the OXML addnext 
            # Use manual bullet formatting since 'List Bullet' style is missing
            current_p = target_p
            for line in lines:
                new_p = doc.add_paragraph(style='normal') # Fallback to normal
                p_fmt = new_p.paragraph_format
                p_fmt.left_indent = Pt(18) # Indent for bullet
                p_fmt.first_line_indent = Pt(-18) # Hanging indent
                
                # Format text
                runner = new_p.add_run("• ") # Manual bullet char
                
                if "**" in line:
                    clean_line = line.replace("* ", "").replace("- ", "")
                    # Split only on the first colon to separate Key from Value
                    parts = clean_line.split(":", 1)
                    if len(parts) > 1:
                        # parts[0] is the key (e.g. "**Data Collection**")
                        # We strip ** and bold it
                        key_text = parts[0].replace("**", "")
                        new_p.add_run(key_text + ":").bold = True
                        
                        # parts[1] is the rest. It might have markdown links, but for now we just add it.
                        # Using existing logic (strip ** just in case? Usually keys have it)
                        new_p.add_run(parts[1])
                    else:
                        # No colon, just bold the whole thing if it has **? 
                        # Or just strip ** and add regular? 
                        # The user example has **Key**: Value.
                        # If just **Text**, let's bold it.
                        new_p.add_run(clean_line.replace("**", ""))
                else:
                    new_p.add_run(line.replace("* ", "").replace("- ", ""))
                
                # Move this new_p to right position
                new_p._element.getparent().remove(new_p._element)
                current_p._element.addnext(new_p._element)
                current_p = new_p # Advance

    insert_lines_after_header(["Section A:"], data["sections"]["A"])
    insert_lines_after_header(["Section B:"], data["sections"]["B"])
    insert_lines_after_header(["Section C:"], data["sections"]["C"])

    # 3. Update Table (Section D)
    if doc.tables:
        table = doc.tables[0] # Assuming first table
        raw_rows = data["table"]
        
        # We need to preserve row 0 (header).
        # We have N data rows.
        # Template has M existing data rows.
        
        # Fill existing rows
        for i, row_data in enumerate(raw_rows):
            # i starts at 0, table row index is i+1 (header is 0)
            table_idx = i + 1
            
            if table_idx < len(table.rows):
                row = table.rows[table_idx]
                row.cells[0].text = row_data[0] # Challenge
                row.cells[1].text = row_data[1] # Approach
                # row.cells[2] is Dropdown. DO NOT TOUCH.
                
                # Optional: try to bold the Challenge text?
                # row.cells[0].paragraphs[0].runs[0].bold = True
            else:
                # Add new row if we run out of template rows
                # This new row won't have the dropdown SDT :(
                # For this specific task, we have 2 items and template has ~3. 
                # We should be fine.
                pass
        
        # Delete unused template rows
        # e.g. if we used rows 1 and 2, and template has 3, 4... delete 3, 4.
        start_del = len(raw_rows) + 1
        while len(table.rows) > start_del:
            # Delete last row
            row = table.rows[-1]
            row._element.getparent().remove(row._element)

    # 4. Footer (Progress Link)
    doc.add_paragraph("_" * 30)
    for line in data["footer"]:
        p = doc.add_paragraph()
        process_line_with_links(p, line)

    doc.save(output_path)
    print(f"Filled Template Saved: {output_path}")

if __name__ == "__main__":
    txt_source = r"c:\Users\dhl\data\thesis\thesis\Thesis_Draft\Updates\2026-02-01_Status_Update_GDOC_READY.txt"
    tpl_source = r"c:\Users\dhl\data\thesis\thesis\Thesis_Draft\Updates\Templates\Weekly or Biweekly Status Update Tempate.docx"
    output = r"c:\Users\dhl\data\thesis\thesis\Thesis_Draft\Updates\2026-02-01_Status_Update.docx"
    
    parsed = process_text_file(txt_source)
    fill_template(tpl_source, output, parsed)
