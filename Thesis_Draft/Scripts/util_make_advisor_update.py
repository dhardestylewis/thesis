from docx import Document
from datetime import date

document = Document()
document.add_heading('Weekly Thesis Status Update', 0)

document.add_paragraph('Name: Daniel Hardesty Lewis')
document.add_paragraph('Project Name: Predicting Zoning Opposition')
document.add_paragraph(f'Date: {date.today().strftime("%b %d, %Y")}\n')

document.add_heading('Section A: What I accomplished in the past week', level=1)
document.add_paragraph('Conducted 5 research interviews with stakeholders out of my target of 10, covering key categories:', style='List Bullet')
document.add_paragraph('City Officials or Planners: Jonathan Tomko, Marla Torrado, Annick Beaudet', style='List 2')
document.add_paragraph('Housing Advocates: Felicity Maxwell', style='List 2')
document.add_paragraph('Residential Developers / Real Estate: Alina Carnahan', style='List 2')
document.add_paragraph('Completed full outline skeleton for all chapters, integrating quantitative modeling with qualitative interviews.', style='List Bullet')
document.add_paragraph('Drafted the Methodology chapter focusing on the complete data cleaning and acquisition pipeline.', style='List Bullet')
document.add_paragraph('Finalized Project One-Pager with AI forecasting hook.', style='List Bullet')
document.add_paragraph('Converted IRB Protocol to PDF with professional formatting.', style='List Bullet')
document.add_paragraph('Standardized EARS files (2019-2025), renamed Bibliography folders, and completed directory hygiene.', style='List Bullet')

document.add_heading("Section B: What I had planned to get done, but didn't get to", level=1)
document.add_paragraph('Interview follow-up (Phase 2) for non-responders.', style='List Bullet')
document.add_paragraph('Sending finalized, personalized invitations to the Columbia reader shortlist.', style='List Bullet')
document.add_paragraph('Training baseline XGBoost/Regression preliminary models.', style='List Bullet')

document.add_heading('Section C: Tasks to complete in the next week', level=1)
document.add_paragraph('Incorporate personalization from Reference PDF for 15 non-responders and execute via fetch_self_outreach.py to secure the remaining 5 interviews.', style='List Bullet')
document.add_paragraph('Send reader outreach invitations.', style='List Bullet')
document.add_paragraph('Schedule mid-February/March advisor check-in meeting.', style='List Bullet')
document.add_paragraph('Train preliminary models on confirmed 2019-2025 longitudinal dataset and draft preliminary results.', style='List Bullet')

document.add_heading('Section D: Challenges/issues/obstacles', level=1)

table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Challenge'
hdr_cells[1].text = 'Approach'
hdr_cells[2].text = 'Do you want to discuss this?'

row_cells = table.add_row().cells
row_cells[0].text = 'Need to locate/confirm layout files for EARS data dictionaries (especially 2018)'
row_cells[1].text = 'Review older archives or request from TCAD directly'
row_cells[2].text = 'No, I am handling it'

row_cells = table.add_row().cells
row_cells[0].text = 'Establishing automated backup for large excluded data files'
row_cells[1].text = 'Automate transfers to G: Drive/S3'
row_cells[2].text = 'No'

document.add_paragraph('\nLet me know if you need anything else! I have also attached the updated outline and methodology chapter drafts to the email.')

document.save('2026-03-02_Status_Update.docx')
