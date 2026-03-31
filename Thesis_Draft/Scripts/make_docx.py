from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

columbia_blue = RGBColor(0, 33, 71)

sections = document.sections
for section in sections:
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

# Custom Heading 0 (Title)
title = document.add_heading(level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Thesis Outline')
run.font.color.rgb = columbia_blue
run.bold = True

subtitle = document.add_paragraph('Predicting Zoning Opposition')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True
subtitle_format = subtitle.paragraph_format
subtitle_format.space_after = Pt(12)

# Custom Heading 1 Style
h1_style = document.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Calibri'
h1_font.size = Pt(12)
h1_font.color.rgb = columbia_blue
h1_font.bold = True
h1_format = h1_style.paragraph_format
h1_format.space_before = Pt(6)
h1_format.space_after = Pt(2)

# Custom List Paragraph Style -> tighten spacing
list_style = document.styles['List Bullet']
list_format = list_style.paragraph_format
list_format.space_before = Pt(0)
list_format.space_after = Pt(0)
list_format.line_spacing = 1.0

document.add_heading('1. Introduction (3-5 pages)', level=1)
document.add_paragraph('☐ 1.1 Context: Rise of zoning opposition and NIMBYism; protest risks for urban planning.', style='List Paragraph')
document.add_paragraph('☐ 1.2 Case Study: Austin\'s housing market, rapid growth, and history of zoning protests.', style='List Paragraph')
document.add_paragraph('☐ 1.3 Research Objectives: Forecasting per-parcel protest risks via DDPM and LogReg.', style='List Paragraph')

document.add_heading('2. Literature Review (5-7 pages)', level=1)
document.add_paragraph('☐ 2.1 Zoning Opposition Dynamics: NIMBYism, homevoter theory, neighborhood resistance.', style='List Paragraph')
document.add_paragraph('☐ 2.2 Quantitative Urban Planning: Predictive modeling applied to urban phenomena.', style='List Paragraph')
document.add_paragraph('☐ 2.3 Generative Causal Models: DDPM and Invariant Causal Prediction in tabular models.', style='List Paragraph')

document.add_heading('3. Mixed-Methods Data and Methodology (8-10 pages) [Completed]', level=1)
document.add_paragraph('☑ 3.1 Quantitative Sources & Panel: Building the balanced property-year panel (282k parcels).', style='List Paragraph')
document.add_paragraph('☑ 3.2 Causal Prediction & Leakage: Preventing temporal leakage across distinct training years.', style='List Paragraph')
document.add_paragraph('☑ 3.3 Qualitative Stakeholder Interviews: Methodology and thematic coding of 10 interviews.', style='List Paragraph')

document.add_heading('4. Predictive Modeling Approach (Quantitative) (5-8 pages) [Completed]', level=1)
document.add_paragraph('☑ 4.1 Baseline Models: Setting up Logistic Regression and XGBoost benchmarks.', style='List Paragraph')
document.add_paragraph('☑ 4.2 Conditional Diffusion (DDPM): Training the generative model for tabular protest risk.', style='List Paragraph')
document.add_paragraph('☑ 4.3 Evaluation Metrics: RMSE, MAE, CRPS, Scenario Std, out-of-sample validation.', style='List Paragraph')

document.add_heading('5. Quantitative Results and Counterfactuals (8-10 pages) [Completed]', level=1)
document.add_paragraph('☑ 5.1 Baseline Performance: Feature importance and predictive accuracy of traditional models.', style='List Paragraph')
document.add_paragraph('☑ 5.2 Generative Models: Evaluating DDPM accuracy and its realistic counterfactuals.', style='List Paragraph')
document.add_paragraph('☑ 5.3 Comparative Analysis & Policy: Using evaluated models to inform proactive rezoning.', style='List Paragraph')

document.add_heading('6. Qualitative Results: Stakeholder Perspectives (8-10 pages)', level=1)
document.add_paragraph('☐ 6.1 Planners and Advocates: How institutional actors view algorithmic planning tools.', style='List Paragraph')
document.add_paragraph('☐ 6.2 Neighborhood Leaders: The tension between predictive efficiency and civic participation.', style='List Paragraph')
document.add_paragraph('☐ 6.3 Democratic Implications: Synthesizing the boundaries of acceptable use.', style='List Paragraph')

document.add_heading('7. Synthesis and Conclusion (4-6 pages)', level=1)
document.add_paragraph('☐ 7.1 Integrating Findings: Reconciling models with stakeholder demands for transparency.', style='List Paragraph')
document.add_paragraph('☐ 7.2 Limitations: ID crosswalk mismatch rates, assumption constraints, model bounding.', style='List Paragraph')
document.add_paragraph('☐ 7.3 Future Work: Scaling to other jurisdictions; NLP-parsed qualitative petition data.', style='List Paragraph')

footer_p = document.add_paragraph('\nTotal Estimated Length: 41-56 pages (excluding Abstract, References, Appendices)')
footer_p.runs[0].italic = True

document.save('thesis_outline.docx')
